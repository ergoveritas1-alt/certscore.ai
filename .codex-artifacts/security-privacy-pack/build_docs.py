from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parent / "outputs"
OUT.mkdir(parents=True, exist_ok=True)

BLUE = "24527A"
DARK_BLUE = "17324D"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
DARK = "202124"
WHITE = "FFFFFF"
AMBER = "FFF4CE"
RED = "FDECEC"
GREEN = "EAF5EE"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_repeat_keep(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def set_font(run, size=11, bold=None, color=DARK, italic=None) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in (
        ("Title", 23, DARK_BLUE, 0, 4),
        ("Subtitle", 13, MID_GRAY, 0, 14),
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_numbering(doc: Document) -> tuple[int, int]:
    numbering = doc.part.numbering_part.element
    existing_abs = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    next_abs = max(existing_abs or [0]) + 1
    next_num = max(existing_num or [0]) + 1

    def create(abstract_id: int, num_id: int, fmt: str, text: str, font: str | None = None):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        lvl.append(lvl_text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        lvl.append(suff)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "720")
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "720")
        ind.set(qn("w:hanging"), "360")
        p_pr.append(ind)
        lvl.append(p_pr)
        if font:
            r_pr = OxmlElement("w:rPr")
            fonts = OxmlElement("w:rFonts")
            fonts.set(qn("w:ascii"), font)
            fonts.set(qn("w:hAnsi"), font)
            r_pr.append(fonts)
            lvl.append(r_pr)
        abstract.append(lvl)
        numbering.append(abstract)
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abs_id = OxmlElement("w:abstractNumId")
        abs_id.set(qn("w:val"), str(abstract_id))
        num.append(abs_id)
        numbering.append(num)

    create(next_abs, next_num, "bullet", "\uf0b7", "Symbol")
    create(next_abs + 1, next_num + 1, "decimal", "%1.")
    return next_num, next_num + 1


def set_num(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.167


def clone_num(doc: Document, source_num_id: int) -> int:
    numbering = doc.part.numbering_part.element
    source = numbering.find(f"{qn('w:num')}[@{qn('w:numId')}='{source_num_id}']")
    if source is None:
        raise ValueError(f"Numbering definition {source_num_id} not found")
    abstract = source.find(qn("w:abstractNumId"))
    abstract_id = abstract.get(qn("w:val"))
    existing = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    new_id = max(existing or [0]) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_id))
    abs_id = OxmlElement("w:abstractNumId")
    abs_id.set(qn("w:val"), abstract_id)
    num.append(abs_id)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    override.append(start_override)
    num.append(override)
    numbering.append(num)
    return new_id


def create_doc(doc_type: str, title: str, subtitle: str) -> tuple[Document, int, int]:
    doc = Document()
    configure_styles(doc)
    bullet_id, number_id = add_numbering(doc)
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run(f"CERTSCORE  |  {doc_type.upper()}")
    set_font(run, size=8.5, color=MID_GRAY, bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Internal governance record  |  Page ")
    set_font(run, size=8.5, color=MID_GRAY)
    add_page_field(footer)

    p = doc.add_paragraph(style="Title")
    p.add_run(title)
    p = doc.add_paragraph(style="Subtitle")
    p.add_run(subtitle)
    metadata = [
        ("Organization", "CertScore (replace with the legal entity name if different)"),
        ("Version", "1.0"),
        ("Status", "Ready for approval - not effective until signed"),
        ("Owner", "Product Owner"),
        ("Effective date", "[enter adoption date]"),
        ("Next review", "12 months after adoption, or sooner after a material change"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{label}: ")
        set_font(r, bold=True)
        r = p.add_run(value)
        set_font(r)
    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(8)
    rule.paragraph_format.space_after = Pt(14)
    p_pr = rule._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), BLUE)
    borders.append(bottom)
    p_pr.append(borders)
    return doc, bullet_id, number_id


def add_bullet(doc, text: str, bullet_id: int, bold_prefix: str | None = None):
    p = doc.add_paragraph()
    set_num(p, bullet_id)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_font(r, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        set_font(r)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_number(doc, text: str, number_id: int):
    p = doc.add_paragraph()
    set_num(p, number_id)
    r = p.add_run(text)
    set_font(r)
    return p


def add_callout(doc, label: str, text: str, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360], 120)
    set_cell_shading(table.cell(0, 0), fill)
    p = table.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{label}: ")
    set_font(r, bold=True, color=DARK_BLUE)
    r = p.add_run(text)
    set_font(r)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[int], header_fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        set_font(r, size=9.5, bold=True, color=DARK_BLUE)
    set_repeat_table_header(table.rows[0])
    prevent_row_split(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_font(r, size=9.3)
    set_table_geometry(table, widths, 120)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_source(doc: Document, label: str, url: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), doc.part.relate_to(url, RT.HYPERLINK, is_external=True))
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    run.append(r_pr)
    text = OxmlElement("w:t")
    text.text = label
    run.append(text)
    hyperlink.append(run)
    p._p.append(hyperlink)


def add_approval(doc):
    doc.add_heading("Approval", level=1)
    p = doc.add_paragraph("This document becomes effective only when the authorized approver completes the fields below. Electronic approval recorded in the evidence repository is acceptable.")
    add_table(doc, ["Approver", "Role", "Approval / signature", "Date"], [["[name]", "Product Owner", "[sign or link to approval record]", "[yyyy-mm-dd]"]], [2100, 1800, 3660, 1800], LIGHT_GRAY)


def save(doc: Document, filename: str):
    props = doc.core_properties
    props.title = filename.replace("_", " ").replace(".docx", "")
    props.subject = "CertScore security and privacy governance"
    props.author = "CertScore"
    props.keywords = "security, privacy, risk, incident response, DPIA"
    doc.save(OUT / filename)


def build_risk_policy():
    doc, bullets, nums = create_doc("Policy", "Information Security Risk Management Policy", "A lightweight process for identifying, treating, and reviewing security risk")
    add_callout(doc, "Purpose", "Establish a proportionate, documented risk-management process suitable for CertScore's current size and AWS-based product environment.")
    doc.add_heading("1. Scope", level=1)
    doc.add_paragraph("This policy applies to CertScore systems, personnel, source code, AWS infrastructure, customer and account information, public-site scan evidence, service providers, endpoints, and business operations. It covers WC01 and WS01 while preserving their separate product and scanner responsibilities.")
    doc.add_heading("2. Policy statements", level=1)
    for t in [
        "The Product Owner is accountable for the security risk-management process and may assign individual risk owners.",
        "Material risks must be recorded in the Security & Privacy Operational Registers workbook, assessed consistently, and assigned an owner and treatment decision.",
        "The risk register is reviewed quarterly and after material architecture, supplier, incident, legal, or threat changes.",
        "Controls and treatments must be proportionate to risk, cost, contractual duties, and legal obligations. Any treatment that may increase recurring or per-request cost follows the repository's cost-change approval rule.",
        "Accepted risks require a written rationale, review date, and approval by the Product Owner. Critical risks cannot be accepted without documented approval.",
        "Evidence supporting completed treatments is linked or described in the register. Sensitive evidence must not be copied into the register when a restricted reference is sufficient.",
    ]:
        add_bullet(doc, t, bullets)
    doc.add_heading("3. Roles", level=1)
    add_table(doc, ["Role", "Minimum responsibility"], [
        ["Product Owner", "Own policy; approve risk acceptance; chair quarterly review; ensure resources and follow-up."],
        ["Risk owner", "Maintain the risk record, implement treatment, retain evidence, and escalate delays or changed exposure."],
        ["Engineering / operations", "Identify technical risks, operate controls, report exceptions and incidents, and support testing."],
        ["All personnel and contractors", "Protect credentials and information; promptly report suspected weaknesses, incidents, or policy exceptions."],
    ], [2100, 7260])
    doc.add_heading("4. Risk assessment method", level=1)
    doc.add_paragraph("Each risk is written as a plausible event affecting an asset or process. Score inherent risk before considering controls, then score residual risk after considering controls.")
    add_table(doc, ["Score", "Likelihood", "Impact"], [
        ["1", "Rare; exceptional circumstances", "Negligible operational, confidentiality, legal, or customer effect"],
        ["2", "Unlikely; could occur", "Limited, quickly recoverable effect"],
        ["3", "Possible; credible occurrence", "Material disruption, disclosure, contractual, or customer effect"],
        ["4", "Likely; expected periodically", "Serious outage, disclosure, regulatory, or reputational effect"],
        ["5", "Almost certain or active", "Severe or sustained harm, major breach, or business-threatening effect"],
    ], [900, 3000, 5460])
    doc.add_paragraph("Risk score = likelihood x impact.")
    add_table(doc, ["Score", "Rating", "Required response"], [
        ["1-4", "Low", "Manage through normal operations; review on schedule."],
        ["5-9", "Moderate", "Assign treatment or acceptance rationale; track to completion."],
        ["10-14", "High", "Product Owner review; treatment plan and target date required."],
        ["15-25", "Critical", "Immediate escalation; begin mitigation or formally pause the affected activity."],
    ], [1200, 1500, 6660])
    doc.add_heading("5. Operating procedure", level=1)
    for t in [
        "Identify: capture the asset/process, threat or failure event, plausible consequence, source, and affected parties.",
        "Assess: assign inherent likelihood and impact, identify existing controls, then assign residual likelihood and impact.",
        "Treat: choose mitigate, avoid, transfer/share, or accept. Record specific actions, owner, target date, and evidence.",
        "Monitor: risk owners update changes and overdue actions. Critical changes are escalated immediately.",
        "Review: the Product Owner conducts and records a quarterly review; closed risks remain retained for at least two years unless a longer contractual or legal period applies.",
    ]:
        add_number(doc, t, nums)
    doc.add_heading("6. Mandatory review triggers", level=1)
    for t in [
        "A confirmed or suspected security/privacy incident, material control failure, or significant vulnerability.",
        "A new production service, high-impact feature, data type, privileged integration, infrastructure pattern, or critical supplier.",
        "A material change to AWS architecture, authentication, encryption, logging, backups, retention, or scanner behavior.",
        "A new or changed legal, regulatory, customer, insurance, or contractual obligation.",
    ]:
        add_bullet(doc, t, bullets)
    doc.add_heading("7. Evidence and exceptions", level=1)
    doc.add_paragraph("Minimum evidence is the current register, dated review record, treatment artifacts, and approval records. Exceptions must state scope, reason, compensating controls, owner, expiry date, and approval. Expired exceptions are treated as open risks.")
    doc.add_heading("8. Review cadence", level=1)
    add_table(doc, ["Activity", "Cadence", "Record"], [
        ["Risk-owner update", "As conditions change; at least quarterly", "Risk Register"],
        ["Formal risk review", "Quarterly", "Review Schedule / meeting record"],
        ["Policy review", "Annually and after material change", "Document history and approval"],
        ["Incident-driven review", "After each material incident", "Incident record and updated risks"],
    ], [3000, 2400, 3960])
    doc.add_heading("9. References", level=1)
    add_source(doc, "NIST Cybersecurity Framework 2.0 Small Business Quick-Start Guide", "https://csrc.nist.gov/pubs/sp/1300/final")
    add_source(doc, "NIST Cybersecurity Framework 2.0", "https://www.nist.gov/cyberframework")
    add_approval(doc)
    save(doc, "01_Information_Security_Risk_Management_Policy.docx")


def build_ir():
    doc, bullets, nums = create_doc("Procedure", "Security Incident Response and Personal Data Breach Procedure", "A single, bounded playbook for technical response, breach assessment, and time-sensitive notification")
    add_callout(doc, "Emergency rule", "If an event may be actively harming systems or people, contain it safely and notify the Incident Lead immediately. Do not wait for complete information.", RED)
    doc.add_heading("1. Purpose and scope", level=1)
    doc.add_paragraph("This procedure applies to suspected or confirmed events affecting CertScore systems, credentials, source code, AWS resources, endpoints, suppliers, customer/account data, scan evidence, or service availability. It covers security incidents and personal data breaches, including events reported by processors or suppliers.")
    doc.add_heading("2. Roles and contact record", level=1)
    add_table(doc, ["Role", "Authority / duty", "Named contact"], [
        ["Incident Lead", "Classifies severity; coordinates containment, investigation, recovery, communications, and closure.", "[name / phone / email]"],
        ["Technical Lead", "Preserves evidence; contains and remediates; maintains technical timeline.", "[name / phone / email]"],
        ["Privacy Lead", "Determines whether personal data is involved; coordinates risk assessment and notification advice.", "[name / phone / email]"],
        ["Communications owner", "Approves accurate customer, partner, insurer, or public communications.", "[name / phone / email]"],
        ["External counsel / adviser", "Provides jurisdiction-specific legal advice where needed.", "[firm / emergency contact]"],
    ], [1700, 5000, 2660])
    add_callout(doc, "Activation", "Any person may activate this procedure by contacting the Incident Lead. If the Incident Lead is unavailable, contact the Product Owner, then the Technical Lead.")
    doc.add_heading("3. What to report", level=1)
    for t in [
        "Suspected unauthorized access, credential exposure, malicious code, data disclosure, alteration, loss, or unavailable data.",
        "Unexpected privileged activity, production compromise, vulnerable dependency exploitation, lost device, phishing, or supplier notification.",
        "Material outage, denial of service, corrupted backups, or control failure with security implications.",
        "Any accidental email, upload, log, support exchange, repository change, or scan artifact exposure involving personal or confidential information.",
    ]:
        add_bullet(doc, t, bullets)
    doc.add_heading("4. Severity", level=1)
    add_table(doc, ["Level", "Examples", "Initial response target"], [
        ["SEV-1 Critical", "Active compromise; material sensitive-data exposure; widespread outage; likely serious harm.", "Immediate; page Incident Lead and Product Owner"],
        ["SEV-2 High", "Confirmed contained compromise; limited data exposure; significant service or supplier event.", "Within 1 hour"],
        ["SEV-3 Moderate", "Suspicious activity or control failure with limited impact and no confirmed compromise.", "Same business day"],
        ["SEV-4 Low", "Minor event, unsuccessful attempt, or policy issue without material impact.", "Within 2 business days"],
    ], [1500, 5160, 2700])
    doc.add_heading("5. Response workflow", level=1)
    for t in [
        "Detect and open a record. Record discovery time, reporter, systems, initial facts, and evidence location. Use UTC for the technical timeline.",
        "Triage and classify. Decide severity, whether personal data may be involved, who must be engaged, and whether the event is continuing.",
        "Contain safely. Revoke exposed credentials, isolate affected resources, block malicious paths, or disable a feature as appropriate. Preserve logs and snapshots before destructive remediation when safe.",
        "Investigate. Establish what happened, attack/failure path, affected systems and data, actors, time period, and confidence. Record facts separately from hypotheses.",
        "Eradicate and recover. Remove the cause, patch or rebuild, validate configuration and credentials, restore service, and monitor for recurrence.",
        "Communicate and notify. Use the breach workflow below and any contractual, insurer, customer, law-enforcement, or sector-specific requirements.",
        "Close and learn. Confirm recovery, document decisions, assign improvements, update the risk register, and hold a lessons-learned review within 10 business days for SEV-1/2.",
    ]:
        add_number(doc, t, nums)
    doc.add_heading("6. Personal data breach workflow", level=1)
    add_callout(doc, "Start the clock", "Record the date and time when CertScore has a reasonable degree of certainty that a security incident compromised personal data. Do not defer this timestamp while waiting for a complete investigation.", AMBER)
    breach_nums = clone_num(doc, nums)
    for t in [
        "Determine role and jurisdiction: controller, processor, or both; affected individuals' locations; lead/relevant supervisory authority; and applicable contracts.",
        "If acting as a processor, notify the relevant controller without undue delay and follow the contract. Do not wait for a controller-level risk determination.",
        "If acting as controller, assess likely risk to individuals using data sensitivity, volume, identifiability, security protections, ease of misuse, affected individuals, and plausible consequences.",
        "Notify the relevant EU/UK authority without undue delay and, where feasible, within 72 hours after awareness when the applicable legal threshold is met. If reporting late, document reasons. Submit information in phases if necessary.",
        "Inform affected individuals without undue delay when the applicable high-risk threshold is met, using clear language and practical protective steps, unless a lawful exception applies.",
        "Record every personal data breach and the notification decision, including non-notification rationale, evidence considered, actions, and approvals.",
    ]:
        add_number(doc, t, breach_nums)
    doc.add_heading("7. 72-hour decision schedule", level=1)
    add_table(doc, ["Elapsed time", "Minimum action"], [
        ["0-4 hours", "Open record; preserve evidence; contain; identify Incident and Privacy Leads; establish awareness time and known jurisdictions."],
        ["By 12 hours", "Initial data/individual impact assessment; controller/processor analysis; contact counsel/adviser for uncertainty or material risk."],
        ["By 24 hours", "Draft notification or documented non-notification rationale; identify authority and required contractual notices."],
        ["By 48 hours", "Decision review and approval; fill known notification fields; plan phased follow-up for missing facts."],
        ["Before 72 hours", "Submit required regulator notification; retain confirmation. If late, document reasons and notify without further delay."],
        ["After submission", "Update authority and affected parties without undue delay; continue containment, investigation, and logging."],
    ], [1800, 7560])
    doc.add_heading("8. Notification information checklist", level=1)
    for t in [
        "Nature of the breach, including confidentiality, integrity, and/or availability impact.",
        "Categories and approximate numbers of affected individuals and personal-data records, where possible.",
        "Privacy/DPO or other contact point.",
        "Likely consequences for individuals.",
        "Measures taken or proposed to address and mitigate adverse effects.",
        "Awareness time, notification time, missing information, planned updates, and any reason for delay.",
    ]:
        add_bullet(doc, t, bullets)
    doc.add_heading("9. Communications and evidence rules", level=1)
    for t in [
        "Only authorized owners communicate externally. Communications must distinguish confirmed facts from estimates and avoid unsupported legal conclusions.",
        "Preserve relevant logs, alerts, tickets, IAM records, configuration, screenshots, vendor communications, and decision records in a restricted evidence location.",
        "Do not place secrets, bearer tokens, raw sensitive URLs, or unnecessary personal data into broad chat, tickets, or the operational register.",
        "Maintain chain-of-custody notes when evidence may support litigation, insurance, law enforcement, or regulatory review.",
    ]:
        add_bullet(doc, t, bullets)
    doc.add_heading("10. Testing and maintenance", level=1)
    doc.add_paragraph("Run a tabletop exercise at least annually and after material changes to personnel or architecture. Test contacts, escalation, access to evidence, containment authority, breach-clock decisions, notification drafting, and restoration. Record gaps, owners, and due dates.")
    doc.add_heading("11. References", level=1)
    add_source(doc, "ICO personal data breach guidance", "https://ico.org.uk/for-organisations/report-a-breach/personal-data-breach/personal-data-breaches-a-guide/")
    add_source(doc, "EDPB Guidelines 9/2022 on personal data breach notification", "https://www.edpb.europa.eu/our-work-tools/our-documents/guidelines/guidelines-92022-personal-data-breach-notification-under_en")
    add_source(doc, "NIST SP 800-61 Rev. 3", "https://csrc.nist.gov/pubs/sp/800/61/r3/final")
    add_approval(doc)
    save(doc, "02_Incident_Response_and_Breach_Notification_Procedure.docx")


def build_privacy():
    doc, bullets, nums = create_doc("Procedure", "Privacy Risk Assessment and DPIA Procedure", "Routine privacy screening, triggered DPIAs, and periodic review")
    add_callout(doc, "Core rule", "Every material change involving personal data receives a documented privacy screening before production. A full DPIA is required when the proposed processing is likely to result in high risk to individuals.")
    doc.add_heading("1. Scope", level=1)
    doc.add_paragraph("This procedure applies to new or materially changed products, scan and report data flows, analytics, account or support functions, AI/model use, vendors, integrations, cookies, retention, sharing, transfers, security controls, and uses of personal data.")
    doc.add_heading("2. Roles", level=1)
    add_table(doc, ["Role", "Responsibility"], [
        ["Product Owner", "Ensures screening occurs; approves risk acceptance and launch decisions; commissions advice where necessary."],
        ["Feature/process owner", "Provides accurate purpose, data-flow, vendor, retention, security, and user-impact details; implements mitigations."],
        ["Privacy reviewer", "Performs or reviews screening/DPIA; tests necessity and proportionality; records conclusion and residual risk."],
        ["Engineering / operations", "Confirms actual architecture, access, logging, retention, deletion, transfer, and security behavior."],
    ], [2200, 7160])
    doc.add_heading("3. When screening is required", level=1)
    for t in [
        "Before launching a new product, feature, processing purpose, personal-data category, tracking technology, or AI/model workflow.",
        "Before onboarding a processor or integration that receives personal data, or materially changing a vendor's purpose, location, subprocessors, or retention.",
        "When processing scale, matching, monitoring, automation, profiling, identifiability, retention, access, disclosure, or geography materially changes.",
        "After a material incident, complaint trend, regulatory change, or evidence that an existing assessment is no longer accurate.",
    ]:
        add_bullet(doc, t, bullets)
    doc.add_heading("4. Screening procedure", level=1)
    for t in [
        "Describe the processing: purpose, data subjects, data fields/categories, sources, recipients, systems, countries, retention, deletion, and role (controller/processor).",
        "Confirm necessity and proportionality: why each data element and recipient is needed; whether a less intrusive design could meet the purpose; how notices, choice, and rights are supported.",
        "Assess security and privacy risks: unauthorized access, misuse, inaccuracy, excessive collection/retention, unexpected use, re-identification, discrimination, exclusion, surveillance, or inability to exercise rights.",
        "Apply the DPIA trigger test. Record each trigger as Yes, No, or Unknown with rationale and evidence.",
        "Decide: no DPIA presently required; full DPIA required before launch; or review required because evidence is incomplete. Unknown does not become No solely because evidence is absent.",
        "Approve, retain, implement mitigations, and set a review date. Re-screen if the implementation changes materially.",
    ]:
        add_number(doc, t, nums)
    doc.add_heading("5. DPIA trigger checklist", level=1)
    add_table(doc, ["Trigger", "Examples / interpretation"], [
        ["Systematic and extensive evaluation", "Profiling or automated assessment that produces legal or similarly significant effects."],
        ["Large-scale sensitive or criminal-offence data", "Special-category, highly sensitive, or offence data processed at material scale."],
        ["Systematic monitoring of public areas", "Persistent or large-scale observation, location, or behavioral tracking."],
        ["Innovative technology plus risk factors", "Novel technology combined with scale, monitoring, vulnerability, matching, or meaningful consequences."],
        ["Data matching or combining", "Datasets combined in ways individuals would not reasonably expect, especially across contexts."],
        ["Vulnerable individuals", "Children, employees, patients, financially vulnerable people, or material power imbalance."],
        ["Denial of service, benefit, or right", "Processing determines access to a service, opportunity, benefit, or exercise of rights."],
        ["Other likely high risk", "Nature, scope, context, and purpose create a credible likelihood of significant harm."],
    ], [3000, 6360])
    add_callout(doc, "Conservative outcome", "If a trigger is plausibly present and cannot be resolved through screening, perform the DPIA or obtain qualified advice before production.", AMBER)
    doc.add_heading("6. Minimum DPIA contents", level=1)
    for t in [
        "A systematic description of processing operations, purposes, roles, data flows, systems, recipients, locations, and retention.",
        "An assessment of necessity and proportionality, including data minimization and less intrusive alternatives.",
        "An assessment of risks to individuals' rights and freedoms, including likelihood, severity, affected groups, and evidence.",
        "Measures addressing each risk, expected residual risk, owner, due date, and evidence of implementation.",
        "Stakeholder, data-subject, DPO, counsel, security, or technical consultation where appropriate.",
        "A signed decision to proceed, modify, defer, or stop. Where high residual risk cannot be mitigated, obtain required supervisory consultation before processing.",
    ]:
        add_bullet(doc, t, bullets)
    doc.add_heading("7. Privacy/DPIA assessment form", level=1)
    add_table(doc, ["Field", "Response"], [
        ["Project / processing activity", "[name and owner]"],
        ["Purpose and expected benefit", "[specific purpose; avoid broad future-use language]"],
        ["Individuals and data", "[categories, sources, approximate volume, sensitivity]"],
        ["Data flow", "[collection -> use -> storage -> recipients -> deletion]"],
        ["Roles and jurisdictions", "[controller/processor; locations and transfers]"],
        ["Necessity / alternatives", "[why needed; lower-impact options considered]"],
        ["Transparency and rights", "[notice, choice, access, deletion, objection, human review]"],
        ["Security and retention", "[access, encryption, logging, backups, deletion schedule]"],
        ["DPIA triggers", "[Yes / No / Unknown for every Section 5 trigger, with rationale]"],
        ["Risks and mitigations", "[scenario, likelihood, impact, action, owner, due date]"],
        ["Decision", "[No DPIA / DPIA required / Review required; approver and date]"],
        ["Review trigger/date", "[material-change triggers and scheduled date]"],
    ], [2700, 6660])
    doc.add_heading("8. Recurring review", level=1)
    doc.add_paragraph("The Product Owner conducts a quarterly review covering new or changed processing, open actions, incidents/complaints, vendor changes, retention/deletion, and overdue assessments. Review completed screenings and DPIAs at least annually and after material changes to the processing or its risks and controls.")
    doc.add_heading("9. Records and retention", level=1)
    doc.add_paragraph("Retain screenings, decisions, approvals, mitigations, evidence references, and review history while processing continues and for at least two years afterward unless a longer period applies. Restrict access and avoid unnecessary personal data.")
    doc.add_heading("10. References", level=1)
    add_source(doc, "ICO Data Protection Impact Assessment guidance", "https://ico.org.uk/for-organisations/uk-gdpr-guidance-and-resources/accountability-and-governance/data-protection-impact-assessments-dpias/")
    add_source(doc, "European Commission: When is a DPIA required?", "https://commission.europa.eu/law/law-topic/data-protection/information-business-and-organisations/obligations/when-data-protection-impact-assessment-dpia-required_en")
    add_approval(doc)
    save(doc, "03_Privacy_Risk_and_DPIA_Procedure.docx")


def build_adoption():
    doc, bullets, nums = create_doc("Assurance record", "Governance Adoption and Assurance Record", "The sign-off and evidence record that turns policy templates into operating processes")
    add_callout(doc, "Current state", "The companion policies and workbook are ready for adoption. Do not represent them as approved, tested, or operating until the corresponding evidence below is completed.", AMBER)
    doc.add_heading("1. Pack contents", level=1)
    add_table(doc, ["Artifact", "Purpose", "Required adoption evidence"], [
        ["Information Security Risk Management Policy", "Formal risk identification, scoring, treatment, acceptance, and quarterly review.", "Signed approval plus completed initial register review"],
        ["Incident Response and Breach Procedure", "Security response and applicable 72-hour personal-data breach workflow.", "Named contacts, signed approval, tabletop record"],
        ["Privacy Risk and DPIA Procedure", "Recurring privacy screening and triggered DPIAs.", "Signed approval plus initial privacy screening review"],
        ["Security & Privacy Operational Registers", "Live risk, incident, breach, screening, schedule, and evidence records.", "Assigned owners, validated entries, dated review"],
    ], [2900, 3700, 2760])
    checklist_heading = doc.add_heading("2. Adoption checklist", level=1)
    checklist_heading.paragraph_format.page_break_before = True
    for t in [
        "Replace 'CertScore' with the correct legal entity name if required.",
        "Name the Incident, Technical, Privacy, and Communications contacts, including out-of-hours contact methods.",
        "Confirm regulator, insurer, legal adviser, hosting, identity, source-control, and critical supplier contacts.",
        "Review and approve the three governance documents; enter effective dates and approval evidence links.",
        "Validate every seeded risk against actual architecture and practices; assign owners and target dates; remove only risks confirmed inapplicable.",
        "Review every seeded privacy activity; correct data fields, purposes, recipients, locations, retention, and role; record the DPIA decision.",
        "Set quarterly review events and an annual tabletop event. Retain attendance and decisions.",
        "Run the tabletop below; record results, corrective actions, owners, and due dates.",
    ]:
        add_bullet(doc, "[ ] " + t, bullets)
    doc.add_heading("3. Initial approval record", level=1)
    add_table(doc, ["Artifact", "Approver", "Decision", "Date", "Evidence link/reference"], [
        ["Risk Management Policy", "[name]", "[approve / changes required]", "[yyyy-mm-dd]", "[location]"],
        ["Incident & Breach Procedure", "[name]", "[approve / changes required]", "[yyyy-mm-dd]", "[location]"],
        ["Privacy & DPIA Procedure", "[name]", "[approve / changes required]", "[yyyy-mm-dd]", "[location]"],
        ["Initial Operational Registers", "[name]", "[validated / changes required]", "[yyyy-mm-dd]", "[location]"],
    ], [2400, 1400, 2200, 1400, 1960])
    doc.add_heading("4. Tabletop exercise", level=1)
    doc.add_paragraph("Suggested scenario: A repository credential is exposed and used to access a production AWS workload. Logs suggest an unknown party accessed a limited set of customer-account records and retained scan artifacts. The exact number of individuals and whether data was downloaded are initially unknown. A critical supplier also reports suspicious activity.")
    doc.add_heading("Exercise objectives", level=2)
    for t in [
        "Activate the plan and confirm authority, contacts, and secure coordination channel.",
        "Contain access without destroying evidence; establish an accurate timeline and awareness time.",
        "Identify controller/processor roles, jurisdictions, likely risks to individuals, and contractual notices.",
        "Make and record a 72-hour notification decision using incomplete information and phased reporting where appropriate.",
        "Recover safely, communicate accurately, and assign lessons-learned actions.",
    ]:
        add_bullet(doc, t, bullets)
    doc.add_heading("Exercise record", level=2)
    add_table(doc, ["Field", "Record"], [
        ["Exercise date / duration", "[yyyy-mm-dd / start-end]"],
        ["Facilitator and participants", "[names and roles]"],
        ["Plan version tested", "[version / effective date]"],
        ["Time to activate / assign severity", "[elapsed time and decision]"],
        ["Containment decisions", "[actions, authority, evidence preserved]"],
        ["Awareness time selected", "[timestamp and rationale]"],
        ["Breach decision", "[notify / do not notify / seek advice; rationale]"],
        ["What worked", "[facts]"],
        ["Gaps", "[facts]"],
        ["Corrective actions", "[action / owner / due date / evidence]"],
        ["Exercise approved by", "[name / date]"],
    ], [2600, 6760])
    doc.add_heading("5. Questionnaire response standard", level=1)
    doc.add_paragraph("Use the following language only after the evidence condition is met. Keep the supporting artifacts available for customer diligence.")
    add_table(doc, ["Question", "Permitted response after adoption", "Evidence condition"], [
        ["Formal information-security risk-management process", "Yes - documented policy, maintained risk register, assigned owners, and quarterly review.", "Signed policy and dated initial/quarterly review"],
        ["Formal documented security incident-response process", "Yes - documented response roles, severity, containment, investigation, recovery, communications, and testing.", "Signed procedure, named contacts, completed tabletop"],
        ["Formal 72-hour breach-reporting process", "Yes - procedure assesses personal-data breaches and supports applicable regulator notification within 72 hours of awareness.", "Signed procedure, contact/decision workflow, completed tabletop"],
        ["Regular DPIA/privacy-risk assessments", "Yes - material changes receive documented privacy screening; high-risk processing receives a DPIA; assessments are reviewed periodically.", "Signed procedure, completed initial screening review, scheduled cadence"],
    ], [2500, 4300, 2560])
    doc.add_heading("6. Important qualification", level=1)
    doc.add_paragraph("A 'Yes' means the process operates as described; it does not mean CertScore is certified, that no risks exist, that every incident is reportable, or that a DPIA is required for every activity. Questionnaire responses must be adjusted if the request uses a different definition, framework, jurisdiction, or audit standard.")
    doc.add_heading("7. Final adoption sign-off", level=1)
    add_table(doc, ["Name", "Role", "Statement", "Signature / approval", "Date"], [["[name]", "Product Owner", "I approve this governance pack for operation and accept responsibility for the scheduled reviews.", "[signature or approval reference]", "[yyyy-mm-dd]"]], [1400, 1400, 3760, 1800, 1000])
    save(doc, "04_Governance_Adoption_and_Assurance_Record.docx")


if __name__ == "__main__":
    build_risk_policy()
    build_ir()
    build_privacy()
    build_adoption()
    print("Created:")
    for path in sorted(OUT.glob("*.docx")):
        print(path)
