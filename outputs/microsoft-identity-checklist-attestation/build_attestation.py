from pathlib import Path
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT_DIR = Path('/Users/benmasek/WC01/outputs/microsoft-identity-checklist-attestation')
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT = OUT_DIR / 'CertScore-Microsoft-Identity-Platform-Checklist-Attestation.docx'
NAVY='0B2545'; BLUE='2E74B5'; LIGHT='E8EEF5'; GRAY='F2F4F7'; INK='172033'; MUTED='5B6573'; GREEN='EAF4EA'; AMBER='FFF8E8'

def run_style(run, size=10, color=INK, bold=False):
    run.font.name='Calibri'; run._element.rPr.rFonts.set(qn('w:ascii'),'Calibri'); run._element.rPr.rFonts.set(qn('w:hAnsi'),'Calibri')
    run.font.size=Pt(size); run.font.color.rgb=RGBColor.from_string(color); run.bold=bold

def ps(p, before=0, after=5, line=1.08, align=None):
    p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(after); p.paragraph_format.line_spacing=line
    if align is not None: p.alignment=align

def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn('w:shd'))
    if shd is None: shd=OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'),fill)

def widths(table, values):
    table.autofit=False; table.alignment=WD_TABLE_ALIGNMENT.LEFT
    tblPr=table._tbl.tblPr; tblW=tblPr.find(qn('w:tblW'))
    if tblW is None: tblW=OxmlElement('w:tblW'); tblPr.append(tblW)
    tblW.set(qn('w:w'),str(sum(values))); tblW.set(qn('w:type'),'dxa')
    ind=tblPr.find(qn('w:tblInd'))
    if ind is None: ind=OxmlElement('w:tblInd'); tblPr.append(ind)
    ind.set(qn('w:w'),'120'); ind.set(qn('w:type'),'dxa')
    grid=table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for value in values:
        col=OxmlElement('w:gridCol'); col.set(qn('w:w'),str(value)); grid.append(col)
    for row in table.rows:
        for i,cell in enumerate(row.cells):
            tcPr=cell._tc.get_or_add_tcPr(); tcW=tcPr.find(qn('w:tcW'))
            if tcW is None: tcW=OxmlElement('w:tcW'); tcPr.append(tcW)
            tcW.set(qn('w:w'),str(values[i])); tcW.set(qn('w:type'),'dxa')
            cell.vertical_alignment=1

def heading(doc, text, level=1):
    p=doc.add_paragraph(style=f'Heading {level}'); ps(p,before=14 if level==1 else 10,after=6); r=p.add_run(text); run_style(r,size=16 if level==1 else 13,color=BLUE,bold=True)

def body(doc, text):
    p=doc.add_paragraph(); ps(p,after=6); r=p.add_run(text); run_style(r)

def bullet(doc, text):
    p=doc.add_paragraph(style='List Bullet'); ps(p,after=3); r=p.add_run(text); run_style(r)

def add_table(doc, headers, rows, col_widths, size=8.3):
    table=doc.add_table(rows=1, cols=len(headers)); widths(table,col_widths)
    trPr=table.rows[0]._tr.get_or_add_trPr(); header=OxmlElement('w:tblHeader'); header.set(qn('w:val'),'true'); trPr.append(header)
    for i,text in enumerate(headers):
        cell=table.rows[0].cells[i]; shade(cell,LIGHT); p=cell.paragraphs[0]; ps(p,after=0,line=1.0); r=p.add_run(text); run_style(r,size=size,color=NAVY,bold=True)
    for row in rows:
        cells=table.add_row().cells
        for i,text in enumerate(row):
            p=cells[i].paragraphs[0]; ps(p,after=0,line=1.02); r=p.add_run(str(text)); run_style(r,size=size)
    widths(table,col_widths); doc.add_paragraph().paragraph_format.space_after=Pt(2)

doc=Document(); sec=doc.sections[0]
sec.page_width=Inches(8.5); sec.page_height=Inches(11); sec.top_margin=Inches(.85); sec.bottom_margin=Inches(.85); sec.left_margin=Inches(.9); sec.right_margin=Inches(.9); sec.header_distance=Inches(.4); sec.footer_distance=Inches(.4)
normal=doc.styles['Normal']; normal.font.name='Calibri'; normal._element.rPr.rFonts.set(qn('w:ascii'),'Calibri'); normal._element.rPr.rFonts.set(qn('w:hAnsi'),'Calibri'); normal.font.size=Pt(10); normal.font.color.rgb=RGBColor.from_string(INK)
for name,size in [('Heading 1',16),('Heading 2',13)]:
    style=doc.styles[name]; style.font.name='Calibri'; style._element.rPr.rFonts.set(qn('w:ascii'),'Calibri'); style._element.rPr.rFonts.set(qn('w:hAnsi'),'Calibri'); style.font.size=Pt(size); style.font.bold=True; style.font.color.rgb=RGBColor.from_string(BLUE)
header=sec.header.paragraphs[0]; ps(header,after=0,line=1.0); r=header.add_run('CertScore.ai  |  Microsoft Identity Platform Checklist'); run_style(r,size=8.5,color=MUTED,bold=True)
footer=sec.footer.paragraphs[0]; ps(footer,after=0,line=1.0,align=WD_ALIGN_PARAGRAPH.RIGHT); r=footer.add_run('Internal attestation worksheet  •  Version 1.0'); run_style(r,size=8,color=MUTED)

p=doc.add_paragraph(); ps(p,after=3); r=p.add_run('MICROSOFT IDENTITY PLATFORM CHECKLIST'); run_style(r,size=22,color=NAVY,bold=True)
p=doc.add_paragraph(); ps(p,after=12); r=p.add_run('Evidence-backed integration review and owner attestation'); run_style(r,size=13,color=MUTED)
add_table(doc,['Control','Value'],[
    ('Application','CertScore.ai Microsoft MCP endpoint'),
    ('Scope','Single-tenant Entra app-only client-credentials integration; AWS-hosted protected MCP API'),
    ('Assessment date','2026-08-23'),
    ('Source','Microsoft identity platform best practices and recommendations'),
    ('Disposition','Technically reviewed; owner attestation required for live/account items'),
],[1800,7560],size=9)

callout=doc.add_table(rows=1,cols=1); widths(callout,[9360]); shade(callout.cell(0,0),AMBER); p=callout.cell(0,0).paragraphs[0]; ps(p,after=0); r=p.add_run('Questionnaire answer: '); run_style(r,size=10.5,color=NAVY,bold=True); r=p.add_run('Select “Yes” after the accountable owner signs this attestation and confirms the owner/action items below. Until then, keep “No.”'); run_style(r,size=10.5)

heading(doc,'1. Review conclusion')
body(doc,'The Microsoft checklist is integration guidance, not a full application audit. This review applies it only to the CertScore Microsoft MCP integration: a single-tenant, service-to-service client-credentials flow. Interactive user, mobile, delegated-consent, and Microsoft Graph-content recommendations are not applicable. The AWS API validates inbound tokens; it does not acquire tokens or call Microsoft Graph.')
body(doc,'The technical evidence supports the scoped controls below. Repository artifacts cannot prove current account ownership, policy acceptance, or live Conditional Access testing, so those items are explicitly assigned to the accountable owner.')

heading(doc,'2. Line-by-line checklist review')
rows=[
('Basics','Read and follow Microsoft platform policies.','Owner attestation','Confirm current platform-policy acceptance and retain the account record.'),
('Ownership','Keep app-management account and registration ownership current.','Owner attestation','Verify live Entra/Partner Center owners; remove unnecessary owners.'),
('Branding','Follow branding guidance; use meaningful name and logo.','Pass','Manifest contains representative CertScore.ai name, website, privacy/terms links, and icons.'),
('Privacy','Provide terms and privacy statement links.','Pass','https://certscore.ai/terms and https://certscore.ai/privacy.'),
('Security','Maintain owned, secure, non-wildcard redirect URIs.','N/A','App-only client credentials; no interactive redirect URI.'),
('Security','Minimize and monitor app-registration owners.','Owner action','Verify the live Entra owner list.'),
('Security','Do not enable implicit grant unless required.','Pass','Client-credentials flow; no implicit grant in documented setup.'),
('Security','Do not use ROPC/password collection.','Pass','No end-user password handling; app-only token flow.'),
('Security','Protect confidential credentials; prefer certs/managed identity; use Key Vault if secret is unavoidable.','Owner exception','Client secret is in Azure Key Vault for certification. Owner must approve the secret-vs-certificate decision and rotation evidence.'),
('Security','Request least-privilege permissions.','Pass','Only the application role `Mcp.Access`; no Graph permissions.'),
('Security','Define API permissions and check expected token permissions.','Pass','Validator checks tenant, issuer, audience, client, token version, and required role.'),
('Implementation','Use modern OAuth 2.0/OpenID Connect.','Pass','Entra v2 OAuth 2.0 client credentials and JWT bearer tokens.'),
('Implementation','Use MSAL when the app acquires tokens.','N/A for AWS API','AWS validates inbound tokens with `jose`; Microsoft certification/runtime owns token acquisition.'),
('Implementation','Parse access tokens only as the protected web API.','Pass','AWS is the resource server and validates inbound tokens before dispatch.'),
('Implementation','Migrate ADAL applications to MSAL.','N/A','No ADAL dependency or legacy ADAL flow.'),
('Implementation','Configure mobile redirects and broker SSO.','N/A','No mobile client in scope.'),
('Implementation','Use a secure token cache per account.','N/A','No interactive user token cache; flow is app-only.'),
('Implementation','Use Graph where required data is available through Graph.','N/A','No Microsoft customer content is required; no Graph permissions.'),
('End-user experience','Provide an understandable consent experience.','Scoped/N/A','No interactive user consent UX; admin consent is recorded for the application role.'),
('End-user experience','Use silent auth before interactive login; avoid prompt=consent.','N/A','No interactive sign-in flow.'),
('End-user experience','Use Graph for optional user-data enrichment.','N/A','No user-data enrichment.'),
('End-user experience','Register required permissions and use incremental consent where applicable.','Scoped pass','One application role with administrator consent; delegated incremental consent is not applicable.'),
('End-user experience','Implement clean single sign-out.','N/A','No interactive user session.'),
('Testing','Test Conditional Access policies.','Owner action','Retain a live workload-identity policy test result before asserting full completion.'),
('Testing','Test all supported account types.','Scoped pass','Single-tenant organizational account type only; consumer, child, mobile, and sovereign types are out of scope.'),
]
add_table(doc,['Area','Checklist practice','Status','Evidence / action'],rows,[1250,3000,1450,3660],size=7.7)

heading(doc,'3. Evidence register')
for item in [
    '/Users/benmasek/WC01/outputs/microsoft-mcp-certification-v2/microsoft_mcp_msft_endpoint_readiness.md',
    '/Users/benmasek/WC01/outputs/microsoft-mcp-certification-v2/azure_setup.md',
    '/Users/benmasek/WC01/apps/mcp/src/microsoft-entra-auth.ts',
    '/Users/benmasek/WC01/outputs/microsoft-mcp-certification-v2/package/manifest.json',
    'https://learn.microsoft.com/en-us/entra/identity-platform/identity-platform-integration-checklist',
]: bullet(doc,item)

heading(doc,'4. Owner attestation and closure')
body(doc,'By signing below, the accountable owner confirms that live Entra/Partner Center owner records are current, Microsoft platform policies are understood and followed, the credential approach and rotation evidence are accepted, and Conditional Access testing has been completed or formally scoped as not applicable.')
add_table(doc,['Attestation item','Record'],[
    ('Accountable owner','____________________________________________'),
    ('Title / organization','____________________________________________'),
    ('Approval date','____________________________________________'),
    ('Credential decision / rotation evidence','____________________________________________'),
    ('Conditional Access test evidence','____________________________________________'),
    ('Next checklist review date','____________________________________________'),
    ('Signature / approval','____________________________________________'),
],[3000,6360],size=9)
final=doc.add_table(rows=1,cols=1); widths(final,[9360]); shade(final.cell(0,0),GREEN); p=final.cell(0,0).paragraphs[0]; ps(p,after=0); r=p.add_run('After signed closure: '); run_style(r,size=10.5,color='1F5E2B',bold=True); r=p.add_run('the organization may answer “Yes” for this scoped Microsoft identity integration.'); run_style(r,size=10.5)

doc.core_properties.title='CertScore.ai Microsoft Identity Platform Checklist Attestation'; doc.core_properties.subject='Evidence-backed checklist review'; doc.core_properties.author='CertScore.ai'; doc.core_properties.comments='Owner signature required for final attestation.'
doc.save(OUT); print(OUT)
