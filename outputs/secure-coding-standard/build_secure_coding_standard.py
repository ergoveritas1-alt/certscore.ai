from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path('/Users/benmasek/WC01/outputs/secure-coding-standard')
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT = OUT_DIR / 'CertScore-Secure-Coding-Standard-OWASP-Top-10.docx'

NAVY = '0B2545'
BLUE = '2E74B5'
LIGHT_BLUE = 'E8EEF5'
LIGHT_GRAY = 'F2F4F7'
INK = '172033'
MUTED = '5B6573'
CAUTION = 'FFF8E8'
RED = '9B1C1C'


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tc_mar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')


def set_table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn('w:tblW'))
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:w'), str(sum(widths)))
    tbl_w.set(qn('w:type'), 'dxa')
    tbl_ind = tbl_pr.find(qn('w:tblInd'))
    if tbl_ind is None:
        tbl_ind = OxmlElement('w:tblInd')
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn('w:w'), '120')
    tbl_ind.set(qn('w:type'), 'dxa')
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement('w:gridCol')
        col.set(qn('w:w'), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn('w:tcW'))
            if tc_w is None:
                tc_w = OxmlElement('w:tcW')
                tc_pr.append(tc_w)
            tc_w.set(qn('w:w'), str(widths[i]))
            tc_w.set(qn('w:type'), 'dxa')
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run(run, size=11, color=INK, bold=False, italic=False, font='Calibri'):
    run.font.name = font
    run._element.rPr.rFonts.set(qn('w:ascii'), font)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def style_paragraph(p, before=0, after=6, line=1.1, align=None):
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if align is not None:
        p.alignment = align


def add_text(doc, text, style=None, size=11, color=INK, bold=False, italic=False, before=0, after=6, line=1.1, align=None):
    p = doc.add_paragraph(style=style)
    style_paragraph(p, before=before, after=after, line=line, align=align)
    r = p.add_run(text)
    set_run(r, size=size, color=color, bold=bold, italic=italic)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    style_paragraph(p, after=4, line=1.1)
    r = p.add_run(text)
    set_run(r)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style='List Number')
    style_paragraph(p, after=5, line=1.1)
    r = p.add_run(text)
    set_run(r)
    # A small normal-paragraph separator keeps LibreOffice from visually
    # joining adjacent numbered paragraphs at a page boundary.
    spacer = doc.add_paragraph()
    style_paragraph(spacer, after=0, line=0.35)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f'Heading {level}')
    r = p.add_run(text)
    if level == 1:
        set_run(r, size=16, color=BLUE, bold=True)
        style_paragraph(p, before=16, after=8, line=1.1)
    elif level == 2:
        set_run(r, size=13, color=BLUE, bold=True)
        style_paragraph(p, before=12, after=6, line=1.1)
    else:
        set_run(r, size=12, color=NAVY, bold=True)
        style_paragraph(p, before=8, after=4, line=1.1)
    return p


def add_callout(doc, label, text, fill=CAUTION, label_color=NAVY):
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [9360])
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    cant_split = OxmlElement('w:cantSplit')
    tr_pr.append(cant_split)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    style_paragraph(p, before=0, after=2, line=1.1)
    r = p.add_run(label + ' ')
    set_run(r, size=10.5, color=label_color, bold=True)
    r = p.add_run(text)
    set_run(r, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths, header_fill=LIGHT_BLUE, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_widths(table, widths)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement('w:tblHeader')
    tbl_header.set(qn('w:val'), 'true')
    tr_pr.append(tbl_header)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        style_paragraph(p, after=0, line=1.0)
        r = p.add_run(header)
        set_run(r, size=font_size, color=NAVY, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            style_paragraph(p, after=0, line=1.05)
            r = p.add_run(str(value))
            set_run(r, size=font_size, color=INK)
    # Ensure all rows use the same explicit cell geometry.
    set_table_widths(table, widths)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Calibri'
normal._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(INK)

for name, size, color, before, after in [
    ('Heading 1', 16, BLUE, 16, 8),
    ('Heading 2', 13, BLUE, 12, 6),
    ('Heading 3', 12, NAVY, 8, 4),
]:
    st = styles[name]
    st.font.name = 'Calibri'
    st._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    st._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.line_spacing = 1.1

header = section.header.paragraphs[0]
style_paragraph(header, after=0, line=1.0)
header.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = header.add_run('CertScore.ai  |  Secure Coding Standard')
set_run(r, size=9, color=MUTED, bold=True)

footer = section.footer.paragraphs[0]
style_paragraph(footer, after=0, line=1.0)
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = footer.add_run('Internal control document  •  Version 1.0')
set_run(r, size=8.5, color=MUTED)

# Memo masthead
add_text(doc, 'SECURE CODING STANDARD', size=23, color=NAVY, bold=True, after=4)
add_text(doc, 'OWASP Top 10-aligned application security requirements', size=14, color=MUTED, after=14)
add_table(doc, ['Control', 'Value'], [
    ('Document owner', 'Product owner / engineering lead'),
    ('Status', 'Draft for organizational adoption'),
    ('Version', '1.0'),
    ('Effective date', 'Upon approval and publication'),
    ('Review cadence', 'At least annually and after material architecture or threat changes'),
    ('Applies to', 'WC01 web, API, MCP, worker, infrastructure-as-code, CI/CD, and supporting operational code'),
], [2100, 7260], header_fill=LIGHT_GRAY, font_size=9.5)

add_callout(doc, 'Questionnaire use:', 'Select “Yes” for OWASP Top 10 secure-coding practices only after the accountable owner adopts this standard, developers are notified, and the required pull-request and scanning evidence is retained.')

add_heading(doc, '1. Purpose and policy statement', 1)
add_text(doc, 'CertScore.ai incorporates security into design, implementation, review, testing, deployment, and maintenance. This standard defines the minimum secure-coding practices used to address common vulnerability classes, including the OWASP Top 10. It is a control standard and operating checklist; it is not a certification or a guarantee that every vulnerability is prevented.')
add_text(doc, 'All production changes must be traceable to a reviewed pull request, pass applicable automated checks, and be evaluated for security impact before deployment. Security defects are prioritized using the vulnerability-management SLAs in the security operations procedure.')

add_heading(doc, '2. Scope and responsibilities', 1)
for text in [
    'Product owner: approves adoption, accepts residual risk, assigns remediation owners, and reviews exceptions.',
    'Change author: applies this standard, records security-relevant design decisions, adds tests, and responds to review findings.',
    'Independent reviewer: checks security impact, tests, authorization, data handling, dependency changes, and rollback implications before approval.',
    'Operations/release owner: confirms CI results, deployment protections, secret handling, and post-deployment monitoring.',
]:
    add_bullet(doc, text)

add_heading(doc, '3. Secure-coding principles', 1)
for text in [
    'Least privilege: grant the minimum identity, permission, network, database, and tool access needed for the task.',
    'Fail closed: reject malformed, unauthenticated, unauthorized, stale, unverifiable, or insufficiently evidenced input and state.',
    'Validate at trust boundaries: parse and validate URLs, headers, tokens, identifiers, JSON, file paths, and external responses before use.',
    'Encode and parameterize: use parameterized database access, context-appropriate output encoding, safe URL handling, and approved serializers.',
    'Minimize data: do not log secrets, tokens, raw credentials, payment-card data, unbounded request/response bodies, or unnecessary personal data.',
    'Make security observable: emit privacy-minimized security events, retain actionable audit records, and alert on material failures.',
    'Use defense in depth: combine application controls, dependency checks, infrastructure controls, secure configuration, and monitoring.',
]:
    add_bullet(doc, text)

add_heading(doc, '4. OWASP Top 10 control mapping', 1)
owasp_rows = [
    ('A01', 'Broken Access Control', 'Enforce authorization on every protected route and tool; bind decisions to authenticated subject, tenant, role, and resource; deny by default; test IDOR and cross-tenant access.'),
    ('A02', 'Cryptographic Failures', 'Use TLS for transport; use approved encryption and managed secret stores; never commit secrets; avoid exposing tokens or sensitive values in logs, URLs, errors, or analytics.'),
    ('A03', 'Injection', 'Use parameterized queries and safe SDK calls; validate and constrain shell, URL, template, HTML, SQL, and JSON inputs; treat scanner content and model output as untrusted data.'),
    ('A04', 'Insecure Design', 'Threat-model new trust boundaries, data flows, authorization changes, external integrations, and high-impact workflows; document abuse cases and safe failure behavior.'),
    ('A05', 'Security Misconfiguration', 'Use secure defaults, explicit allowlists, hardened headers, private storage, least-privilege IAM, disabled debug paths, dependency pinning, and configuration review.'),
    ('A06', 'Vulnerable and Outdated Components', 'Run dependency and infrastructure vulnerability scans; review advisories; patch according to severity SLAs; record mitigations and risk acceptance when a fix is not immediately available.'),
    ('A07', 'Identification and Authentication Failures', 'Use approved authentication flows; validate issuer, audience, tenant, signature, expiry, token type, and required roles; protect sessions and account recovery; never accept role-less protected requests.'),
    ('A08', 'Software and Data Integrity Failures', 'Require reviewed changes, protected branches, reproducible builds where feasible, signed or immutable deployment artifacts, dependency integrity checks, and verification of retained evidence.'),
    ('A09', 'Security Logging and Monitoring Failures', 'Log security-relevant outcomes without secrets; review logs on a defined cadence; alert for authentication, authorization, availability, and integrity anomalies; preserve incident evidence.'),
    ('A10', 'Server-Side Request Forgery', 'Allow outbound requests only to validated, intended destinations; block private/link-local metadata targets; enforce scheme, DNS, redirect, timeout, response-size, and content-type limits.'),
]
add_table(doc, ['ID', 'Vulnerability class', 'Minimum required practice'], owasp_rows, [650, 2250, 6460], header_fill=LIGHT_BLUE, font_size=8.8)

add_heading(doc, '5. Required development workflow', 1)
for i, text in enumerate([
    'Design review: identify trust boundaries, identities, sensitive data, external services, authorization decisions, and abuse cases before implementation.',
    'Implementation: use approved libraries and patterns; validate untrusted input; apply least privilege; keep secrets and sensitive values out of code and logs.',
    'Automated checks: run unit/integration security tests, dependency checks, infrastructure-as-code checks, lint/type checks, and relevant static analysis before merge.',
    'Independent review: a reviewer other than the author verifies security impact, access control, input handling, data exposure, tests, and rollback implications.',
    'Release: deploy only from the protected production workflow after required checks and approvals pass; use immutable image or artifact references where available.',
    'Post-release: monitor security events and errors, investigate anomalies, and record corrective actions or risk acceptance.',
]):
    add_number(doc, text)

add_heading(doc, '6. Pull-request security checklist', 1)
for text in [
    '☐ Does the change alter authentication, authorization, tenant isolation, secrets, payment, privacy, or external network behavior?',
    '☐ Are all new inputs validated for type, length, format, encoding, and allowed values?',
    '☐ Are database, shell, template, URL, HTML, and deserialization operations parameterized or safely constrained?',
    '☐ Are permissions, roles, IAM actions, API scopes, and network destinations least-privilege and deny-by-default?',
    '☐ Could logs, errors, telemetry, screenshots, artifacts, or model prompts expose secrets or personal data?',
    '☐ Are negative tests included for unauthorized, malformed, cross-tenant, stale, expired, and oversized inputs?',
    '☐ Are dependencies, containers, Terraform, and runtime configuration checked for known vulnerabilities or insecure defaults?',
    '☐ Is a reviewer other than the author assigned, and are findings resolved or documented as accepted risk?',
]:
    add_bullet(doc, text)

add_heading(doc, '7. Testing and evidence requirements', 1)
add_text(doc, 'The following records must be retained for production changes when applicable:')
for text in [
    'Pull request URL, author, reviewer, approval, changed scope, and security-impact assessment.',
    'Test results covering authentication, authorization, input validation, negative paths, tenant isolation, and sensitive-data handling.',
    'Dependency, container, and infrastructure scan results, including remediation issue or documented risk acceptance for unresolved findings.',
    'Deployment artifact identifier, environment, approval, and rollback reference.',
    'Security incident, alert, or post-release review record when the change affects a material security control.',
]:
    add_bullet(doc, text)

add_heading(doc, '8. Vulnerability handling and exceptions', 1)
add_text(doc, 'Security vulnerabilities are triaged and remediated under the production patching SLAs: critical or known exploited issues within 72 hours, high within 14 calendar days, medium within 30 calendar days, and low within 90 calendar days. A mitigation may replace a patch only when it is documented and demonstrably reduces exposure.')
add_text(doc, 'An exception requires: the affected control, reason, scope, risk assessment, compensating control, named owner, expiration/review date, and product-owner approval. Exceptions do not authorize logging secrets, bypassing tenant isolation, accepting unauthenticated protected access, or weakening evidence integrity.')

add_heading(doc, '9. Adoption and review', 1)
add_text(doc, 'This standard becomes an operating control when the accountable owner signs below, publishes it to the engineering team, and links it from the production change workflow. Review it at least annually, after a material security incident, and before a material change to authentication, authorization, data processing, external integrations, or deployment architecture.')
add_table(doc, ['Adoption item', 'Record'], [
    ('Accountable owner', '____________________________________________'),
    ('Approval date', '____________________________________________'),
    ('Publication location', '____________________________________________'),
    ('Next review date', '____________________________________________'),
    ('Evidence location', '____________________________________________'),
], [2400, 6960], header_fill=LIGHT_GRAY, font_size=9.5)

add_callout(doc, 'Questionnaire answer:', 'After adoption and evidence retention, the organization may answer “Yes” to: “Do secure coding practices take into account common vulnerability classes such as OWASP Top 10?” Before adoption, answer “No.”', fill='EAF4EA', label_color='1F5E2B')

add_heading(doc, 'Appendix A. References', 1)
for text in [
    'OWASP Top 10 (2021): https://owasp.org/Top10/2021/',
    'NIST Secure Software Development Framework (SSDF): https://csrc.nist.gov/Projects/ssdf',
    'CertScore security operations procedure: /Users/benmasek/WC01/docs/security-operations.md',
    'CertScore vulnerability register: /Users/benmasek/WC01/docs/security-vulnerability-register.md',
]:
    add_bullet(doc, text)

doc.core_properties.title = 'CertScore.ai Secure Coding Standard - OWASP Top 10'
doc.core_properties.subject = 'Secure coding controls and OWASP Top 10 questionnaire evidence'
doc.core_properties.author = 'CertScore.ai'
doc.core_properties.comments = 'Draft for accountable-owner adoption; not a certification.'
doc.save(OUT)
print(OUT)
